"""Resume text extraction from PDF, DOCX, and TXT files."""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Union

import pdfplumber
from docx import Document

FileInput = Union[str, Path, bytes, io.BytesIO]


def _read_bytes(source: FileInput) -> tuple[bytes, str]:
    """Return (file_bytes, filename) for any supported input form."""
    if isinstance(source, (str, Path)):
        path = Path(source)
        return path.read_bytes(), path.name
    if isinstance(source, bytes):
        return source, "uploaded.bin"
    # Streamlit UploadedFile / BytesIO duck-type
    name = getattr(source, "name", "uploaded.bin")
    data = source.read()
    # Reset cursor for any subsequent reads
    if hasattr(source, "seek"):
        source.seek(0)
    return data, name


def _extract_pdf(data: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text:
                text_parts.append(text)
    return "\n".join(text_parts)


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Capture table content too — common in resume templates.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _normalize(text: str) -> str:
    """Collapse whitespace, strip control chars, keep punctuation for NLP."""
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[\t\x0b\x0c]+", " ", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text(source: FileInput, filename: str | None = None) -> str:
    """Extract clean text from a PDF, DOCX, or TXT resume.

    Args:
        source: File path, bytes, or file-like object (e.g. Streamlit upload).
        filename: Optional override; used to detect file type by extension.

    Returns:
        Normalized text content. Empty string if extraction fails.
    """
    data, detected_name = _read_bytes(source)
    name = (filename or detected_name).lower()

    try:
        if name.endswith(".pdf"):
            text = _extract_pdf(data)
        elif name.endswith(".docx"):
            text = _extract_docx(data)
        elif name.endswith((".txt", ".md")):
            text = _extract_txt(data)
        else:
            # Best-effort fallback: try PDF, then DOCX, then TXT.
            try:
                text = _extract_pdf(data)
            except Exception:
                try:
                    text = _extract_docx(data)
                except Exception:
                    text = _extract_txt(data)
    except Exception:
        return ""

    return _normalize(text)
